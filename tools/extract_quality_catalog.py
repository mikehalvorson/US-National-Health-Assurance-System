"""Extract the controlled quality and cost parameter catalogs for the dashboard.

The framework DOCX is the controlled, machine-readable companion to the
published PDF.  This script keeps the dashboard catalog reproducible instead
of maintaining hundreds of records by hand.
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "National_Health_Assurance_Framework_v2.0.0_FINAL.docx"
OUTPUT = ROOT / "docs" / "js" / "qualitydata.js"


CONCEPTS = [
    "Coverage, affordability & eligibility",
    "Access, routing & unit network",
    "Quality, safety & patient experience",
    "Equity, rights & legitimacy",
    "Workforce & care delivery",
    "Claims, payments & financing",
    "Hospitals & regional delivery",
    "Medicines, devices & supply",
    "Expanded benefits & public health",
    "Data, cybersecurity & AI",
    "Governance, transition & continuity",
    "Research, innovation & training",
    "Population, macroeconomics & offsets",
]


TPP_CONCEPTS = {
    "1": CONCEPTS[0],
    "2": CONCEPTS[5],
    "3": CONCEPTS[7],
    "4": CONCEPTS[7],
    "5": CONCEPTS[6],
    "6": CONCEPTS[1],
    "7": CONCEPTS[1],
    "8": CONCEPTS[4],
    "9": CONCEPTS[8],
    "10": CONCEPTS[9],
    "11": CONCEPTS[9],
    "12": CONCEPTS[10],
    "13": CONCEPTS[11],
    "EMP": CONCEPTS[10],
    "LEG": CONCEPTS[3],
    "USE": CONCEPTS[3],
    "TRIB": CONCEPTS[3],
    "REG": CONCEPTS[10],
    "FORM": CONCEPTS[10],
}


CP_CONCEPTS = {
    "CP-TOT": CONCEPTS[12],
    "CP-POP": CONCEPTS[12],
    "CP-CLM": CONCEPTS[5],
    "CP-HOSP": CONCEPTS[6],
    "CP-CLIN": CONCEPTS[4],
    "CP-UNIT": CONCEPTS[1],
    "CP-LTC": CONCEPTS[8],
    "CP-RX": CONCEPTS[7],
    "CP-DX": CONCEPTS[7],
    "CP-BH": CONCEPTS[8],
    "CP-DVH": CONCEPTS[8],
    "CP-EMS": CONCEPTS[8],
    "CP-PH": CONCEPTS[8],
    "CP-IT": CONCEPTS[9],
    "CP-GOV": CONCEPTS[10],
    "CP-RD": CONCEPTS[11],
    "CP-EDU": CONCEPTS[11],
    "CP-TRN": CONCEPTS[10],
    "CP-FIN": CONCEPTS[5],
    "CP-OFF": CONCEPTS[12],
}


def clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def table_rows(table) -> list[dict[str, str]]:
    headers = [clean(cell.text) for cell in table.rows[0].cells]
    return [
        {header: clean(cell.text) for header, cell in zip(headers, row.cells)}
        for row in table.rows[1:]
    ]


def kpp_concept(domain: str) -> str:
    key = domain.lower()
    if key.startswith("coverage"):
        return CONCEPTS[0]
    if key.startswith("access"):
        return CONCEPTS[1]
    if key.startswith("quality") or key.startswith("safety"):
        return CONCEPTS[2]
    if key.startswith("equity") or key.startswith("legitimacy"):
        return CONCEPTS[3]
    if key.startswith("workforce"):
        return CONCEPTS[4]
    if key.startswith("cost"):
        return CONCEPTS[5]
    if key.startswith("transition"):
        return CONCEPTS[10]
    return CONCEPTS[10]


def tpp_key(identifier: str) -> str:
    suffix = identifier.removeprefix("TPP-")
    match = re.match(r"(\d+|[A-Z]+)", suffix)
    if not match:
        raise ValueError(f"Unrecognized TPP identifier: {identifier}")
    return match.group(1)


def source_rows(document: Document) -> tuple[list[dict], list[dict], list[dict]]:
    kpps = []
    for row in table_rows(document.tables[114]):
        kpps.append(
            {
                "id": row["ID"],
                "type": "KPP",
                "name": row["Metric"],
                "concept": kpp_concept(row["Domain / trace"]),
                "where": row["Domain / trace"],
                "target": row["Source target"],
                "calculation": row["Calculation contract"],
                "datasets": row["Datasets"],
                "ownerVerifier": row["Owner / verifier"],
                "status": row["Status"],
                "unit": "",
                "modelRole": "performance outcome",
                "temporal": "maturity target",
                "unitStatus": "",
                "family": "",
            }
        )

    tpps = []
    for row in table_rows(document.tables[116]):
        identifier = row["ID"]
        concept = TPP_CONCEPTS[tpp_key(identifier)]
        tpps.append(
            {
                "id": identifier,
                "type": "TPP",
                "name": row["Metric"],
                "concept": concept,
                "where": concept,
                "target": row["Source target"],
                "calculation": row["Calculation / unit"],
                "datasets": row["Datasets"],
                "ownerVerifier": row["Owner / verifier"],
                "status": row["Status"],
                "unit": "",
                "modelRole": "technical performance",
                "temporal": "maturity target",
                "unitStatus": "",
                "family": "",
            }
        )

    family_rows = table_rows(document.tables[120])
    family_by_id = {row["Family"]: row for row in family_rows}
    cps = []
    for table_index, family_row in zip(range(121, 141), family_rows):
        family = family_row["Family"]
        expected = int(family_row["Records"])
        rows = table_rows(document.tables[table_index])
        if len(rows) != expected:
            raise ValueError(
                f"{family}: expected {expected} records, extracted {len(rows)}"
            )
        for row in rows:
            definition = row["Definition"]
            name = definition.split(" — ", 1)[0]
            cps.append(
                {
                    "id": row["ID"],
                    "type": "CP",
                    "name": name,
                    "concept": CP_CONCEPTS[family],
                    "where": family_row["Domain"],
                    "target": row["Value/source status"],
                    "calculation": definition,
                    "datasets": "",
                    "ownerVerifier": family_row["Accountable owner(s)"],
                    "status": row["Value/source status"],
                    "unit": row["Canonical unit"],
                    "modelRole": row["Model role"],
                    "temporal": row["Temporal"],
                    "unitStatus": row["Unit status"],
                    "family": family,
                }
            )
    return kpps, tpps, cps


def phase_records(document: Document) -> list[dict]:
    records = []
    for row in table_rows(document.tables[483]):
        records.append(
            {
                "id": row["ID"].replace("PH-", ""),
                "anchor": row["Anchor"],
                "purpose": row["Purpose"],
                "scope": row["Controlled products / scope"],
                "exitEvidence": row["Exit evidence"],
            }
        )
    return records


def gate_records(document: Document) -> list[dict]:
    return [
        {
            "id": row["ID"].replace("PH-", ""),
            "name": row["Gate"],
            "decision": row["Decision point"],
            "floor": row["Blocking floor / condition"],
            "evidence": row["Evidence / trace"],
            "fallback": row["Fallback / repair"],
        }
        for row in table_rows(document.tables[485])
    ]


def target_overrides(document: Document) -> dict[str, list[dict]]:
    overrides: dict[str, list[dict]] = {}
    for row in table_rows(document.tables[486]):
        match = re.match(r"(KPP|TPP)-[A-Z0-9.]+", row["Mature target"])
        if not match:
            raise ValueError(f"Cannot identify parameter in {row['Mature target']}")
        identifier = match.group(0)
        phase = {
            "PH-G1": "P4",
            "PH-G2": "P6",
            "PH-G3": "P7",
            "PH-G4": "P8",
            "PH-G5": "P8",
        }[row["Gate"]]
        overrides.setdefault(identifier, []).append(
            {
                "phase": phase,
                "gate": row["Gate"].replace("PH-", ""),
                "value": row["Progression floor"],
                "kind": "progression floor",
                "interpretation": row["Interpretation"],
            }
        )

    # The roadmap establishes one additional numeric phase milestone before G2.
    overrides.setdefault("KPP-B7", []).insert(
        0,
        {
            "phase": "P5",
            "gate": "",
            "value": ">=65% by phase end",
            "kind": "phase milestone",
            "interpretation": "P5 delivery-scale milestone from the controlled phase plan.",
        },
    )
    return overrides


def build_catalog() -> dict:
    document = Document(SOURCE)
    kpps, tpps, cps = source_rows(document)
    overrides = target_overrides(document)
    parameters = kpps + tpps + cps

    for parameter in parameters:
        if parameter["type"] in {"KPP", "TPP"}:
            rollout = list(overrides.get(parameter["id"], []))
            rollout.append(
                {
                    "phase": "P8",
                    "gate": "",
                    "value": parameter["target"],
                    "kind": "maturity target",
                    "interpretation": "Controlled source target assessed at maturity.",
                }
            )
            parameter["rollout"] = rollout
            parameter["phaseNote"] = (
                "No numeric interim target is specified unless an earlier phase "
                "milestone or gate floor appears below."
            )
        else:
            parameter["rollout"] = []
            parameter["phaseNote"] = (
                "Cost/model parameter, not a phase performance target. "
                "The framework does not assign phase-specific values."
            )

    identifiers = [item["id"] for item in parameters]
    if len(identifiers) != len(set(identifiers)):
        raise ValueError("Duplicate parameter identifier found")
    if (len(kpps), len(tpps), len(cps), len(parameters)) != (41, 79, 310, 430):
        raise ValueError(
            "Unexpected catalog counts: "
            f"{len(kpps)} KPP, {len(tpps)} TPP, {len(cps)} CP"
        )

    family_rows = table_rows(document.tables[120])
    return {
        "source": "National Health Assurance Framework v2.0.0 FINAL",
        "counts": {
            "KPP": len(kpps),
            "TPP": len(tpps),
            "CP": len(cps),
            "total": len(parameters),
        },
        "concepts": CONCEPTS,
        "cpFamilies": [
            {
                "id": row["Family"],
                "domain": row["Domain"],
                "records": int(row["Records"]),
                "owners": row["Accountable owner(s)"],
                "sensitivity": row["Primary sensitivity group"],
            }
            for row in family_rows
        ],
        "phases": phase_records(document),
        "gates": gate_records(document),
        "parameters": parameters,
    }


def main() -> None:
    catalog = build_catalog()
    payload = json.dumps(catalog, ensure_ascii=False, separators=(",", ":"))
    OUTPUT.write_text(
        "window.NHA_QUALITY_DATA=" + payload + ";\n",
        encoding="utf-8",
    )
    counts = catalog["counts"]
    print(
        f"Wrote {OUTPUT.relative_to(ROOT)}: "
        f"{counts['KPP']} KPP + {counts['TPP']} TPP + "
        f"{counts['CP']} CP = {counts['total']} records"
    )


if __name__ == "__main__":
    main()
